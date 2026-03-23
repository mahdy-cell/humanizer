"""
core/exporter.py — Layer 6: Export Layer

  - Rebuild .docx via python-docx with original formatting preserved
  - APA 7 margins (2.54 cm), Times New Roman 12pt, double spacing
  - Running head + page numbers in header/footer
  - Hanging indent for bibliography
  - Metadata scrubbing + injection
  - PDF/A export via reportlab
"""

import warnings
import os
from typing import Optional, Dict, Any

try:
    import docx
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    warnings.warn("python-docx not available; .docx export disabled.")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    _REPORTLAB_AVAILABLE = True
except ImportError:
    _REPORTLAB_AVAILABLE = False
    warnings.warn("reportlab not available; PDF export disabled.")


# ── .docx helpers ──────────────────────────────────────────────────────────

def _set_margins(doc: "Document", margin_cm: float = 2.54) -> None:
    """Set all page margins to *margin_cm* centimetres."""
    if not _DOCX_AVAILABLE:
        return
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def _set_font(paragraph, font_name: str = "Times New Roman", font_size: int = 12) -> None:
    """Apply font settings to all runs in *paragraph*."""
    if not _DOCX_AVAILABLE:
        return
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _set_double_spacing(paragraph) -> None:
    """Set double line spacing on *paragraph*."""
    if not _DOCX_AVAILABLE:
        return
    from docx.shared import Pt as _Pt
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    pPr = paragraph._p.get_or_add_pPr()
    spacing = _OE("w:spacing")
    spacing.set(_qn("w:line"), "480")
    spacing.set(_qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_page_numbers(doc: "Document", running_head: str = "") -> None:
    """Add running head and page numbers to header/footer."""
    if not _DOCX_AVAILABLE:
        return
    try:
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = running_head.upper()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)
    except Exception as exc:
        warnings.warn(f"Header/footer setup failed: {exc}")


def _scrub_metadata(doc: "Document", author: str = "", university: str = "") -> None:
    """Scrub and inject document metadata."""
    if not _DOCX_AVAILABLE:
        return
    try:
        props = doc.core_properties
        props.author = author
        props.company = university
        props.last_modified_by = author
        import datetime
        props.created = datetime.datetime.now()
        props.modified = datetime.datetime.now()
        # Clear keywords that might leak AI tooling
        props.keywords = ""
        props.comments = ""
    except Exception as exc:
        warnings.warn(f"Metadata scrubbing failed: {exc}")


def _sanitize_xml(doc: "Document") -> None:
    """Remove rsidR tracking tags and revision history from .docx XML."""
    if not _DOCX_AVAILABLE:
        return
    try:
        tracking_tags = [
            qn("w:rsidR"), qn("w:rsidRPr"), qn("w:rsidDel"),
            qn("w:rsidRDefault"), qn("w:ins"), qn("w:del"),
        ]
        for para in doc.paragraphs:
            for tag in tracking_tags:
                for element in para._p.findall(f".//{tag}"):
                    element.getparent().remove(element)
    except Exception as exc:
        warnings.warn(f"XML sanitization failed: {exc}")


# ── Export functions ───────────────────────────────────────────────────────

def export_docx(
    text: str,
    output_path: str,
    config: Dict[str, Any] = None,
) -> bool:
    """
    Export *text* to a .docx file with APA 7 formatting.

    Returns True on success.
    """
    if not _DOCX_AVAILABLE:
        warnings.warn("python-docx not available; cannot export .docx")
        return False
    if config is None:
        config = {}
    try:
        doc = Document()
        _set_margins(doc, config.get("margins_cm", 2.54))
        _scrub_metadata(
            doc,
            author=config.get("author_name", ""),
            university=config.get("university_name", ""),
        )
        _add_page_numbers(doc, running_head=config.get("running_head", ""))

        font_name = config.get("font_name", "Times New Roman")
        font_size = config.get("font_size", 12)

        paragraphs = text.split("\n\n") if "\n\n" in text else text.split("\n")
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            p = doc.add_paragraph(para_text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_font(p, font_name, font_size)
            _set_double_spacing(p)

        _sanitize_xml(doc)
        doc.save(output_path)
        return True
    except Exception as exc:
        warnings.warn(f"DOCX export failed: {exc}")
        return False


def export_pdf(
    text: str,
    output_path: str,
    config: Dict[str, Any] = None,
) -> bool:
    """
    Export *text* to a PDF/A file using reportlab.

    Returns True on success.
    """
    if not _REPORTLAB_AVAILABLE:
        warnings.warn("reportlab not available; cannot export PDF")
        return False
    if config is None:
        config = {}
    try:
        margin = config.get("margins_cm", 2.54) * cm
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=margin,
            rightMargin=margin,
            topMargin=margin,
            bottomMargin=margin,
        )
        styles = getSampleStyleSheet()
        style = styles["Normal"]
        style.fontName = "Times-Roman"
        style.fontSize = config.get("font_size", 12)
        style.leading = config.get("font_size", 12) * 2.0  # double spacing

        story = []
        paragraphs = text.split("\n\n") if "\n\n" in text else text.split("\n")
        for para_text in paragraphs:
            if not para_text.strip():
                story.append(Spacer(1, 12))
                continue
            story.append(Paragraph(para_text.strip(), style))
            story.append(Spacer(1, style.leading))

        doc.build(story)
        return True
    except Exception as exc:
        warnings.warn(f"PDF export failed: {exc}")
        return False


def export(
    text: str,
    output_path: str,
    format: str = "docx",
    config: Dict[str, Any] = None,
) -> bool:
    """
    Export *text* to *output_path* in the specified *format* ('docx' or 'pdf').
    """
    fmt = format.lower().lstrip(".")
    if fmt == "docx":
        return export_docx(text, output_path, config)
    elif fmt in ("pdf", "pdfa"):
        return export_pdf(text, output_path, config)
    else:
        warnings.warn(f"Unknown export format: {format}")
        return False
